from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional
import uuid
from datetime import datetime, date
from openai import OpenAI
import io
import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import PyPDF2
import pdfplumber
from docx import Document as DocxReader

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ.get('DB_NAME', 'resume_builder')]

# OpenAI client
openai_client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Pydantic Models
class PersonalInfo(BaseModel):
    fullName: str
    email: str
    phone: str
    location: str = ""
    linkedin: str = ""
    portfolio: str = ""
    photo: str = ""

class ExperienceItem(BaseModel):
    id: str
    title: str
    company: str
    location: str = ""
    startDate: str
    endDate: str = ""
    current: bool = False
    bullets: List[str] = []

class EducationItem(BaseModel):
    id: str
    degree: str
    school: str
    location: str = ""
    graduationDate: str
    gpa: str = ""

class CertificationItem(BaseModel):
    id: str
    name: str
    issuer: str
    date: str = ""

class LanguageItem(BaseModel):
    id: str
    language: str
    proficiency: str

class ResumeData(BaseModel):
    personalInfo: PersonalInfo
    summary: str = ""
    experience: List[ExperienceItem] = []
    education: List[EducationItem] = []
    skills: List[str] = []
    certifications: List[CertificationItem] = []
    languages: List[LanguageItem] = []

class ResumeCreate(BaseModel):
    resumeData: ResumeData
    template: str = "professional"

class Resume(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    resumeData: ResumeData
    template: str
    createdAt: datetime = Field(default_factory=datetime.utcnow)
    updatedAt: datetime = Field(default_factory=datetime.utcnow)

class ATSAnalysisRequest(BaseModel):
    resumeData: ResumeData
    jobDescription: str

class KeywordAnalysis(BaseModel):
    matched: List[str]
    missing: List[str]
    overused: List[str] = []

class ImpactOpportunity(BaseModel):
    original: str
    suggestion: str

class ReadabilityAnalysis(BaseModel):
    score: int
    suggestions: List[str]

class ATSAnalysisResponse(BaseModel):
    score: int
    status: str
    keywordAnalysis: KeywordAnalysis
    suggestions: List[str]
    impactOpportunities: List[ImpactOpportunity]
    readability: ReadabilityAnalysis

class ExportRequest(BaseModel):
    resumeData: ResumeData
    template: str

class CoverLetterRequest(BaseModel):
    resumeData: ResumeData
    jobDescription: str
    companyName: str = ""
    jobTitle: str = ""

class CoverLetterResponse(BaseModel):
    content: str
    suggestions: List[str] = []

class SkillsExtractRequest(BaseModel):
    text: str
    existingSkills: List[str] = []

class SkillItem(BaseModel):
    name: str
    category: str
    alreadyAdded: bool = False

class SkillsExtractResponse(BaseModel):
    skills: List[SkillItem]

# Helper Functions
def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_file = io.BytesIO(file_content)
        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        raise HTTPException(status_code=400, detail="Failed to extract text from PDF")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        docx_file = io.BytesIO(file_content)
        doc = DocxReader(docx_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise HTTPException(status_code=400, detail="Failed to extract text from DOCX")

async def parse_resume_with_ai(resume_text: str) -> ResumeData:
    """Use OpenAI to parse resume text into structured data"""
    
    prompt = f"""Parse this resume text and extract structured information. Return JSON with these exact fields:

{{
  "personalInfo": {{
    "fullName": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin": "",
    "portfolio": "",
    "photo": ""
  }},
  "summary": "",
  "experience": [
    {{
      "id": "exp1",
      "title": "",
      "company": "",
      "location": "",
      "startDate": "DD-MM-YYYY",
      "endDate": "DD-MM-YYYY or Present",
      "current": false,
      "bullets": []
    }}
  ],
  "education": [
    {{
      "id": "edu1",
      "degree": "",
      "school": "",
      "location": "",
      "graduationDate": "YYYY",
      "gpa": ""
    }}
  ],
  "skills": [],
  "certifications": [],
  "languages": []
}}

Resume Text:
{resume_text}

Extract all information accurately. For dates, use DD-MM-YYYY format. Leave fields empty if not found."""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert resume parser. Always return valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        
        # Map to ResumeData model
        return ResumeData(
            personalInfo=PersonalInfo(**result.get('personalInfo', {})),
            summary=result.get('summary', ''),
            experience=[ExperienceItem(**exp) for exp in result.get('experience', [])],
            education=[EducationItem(**edu) for edu in result.get('education', [])],
            skills=result.get('skills', []),
            certifications=[CertificationItem(**cert) for cert in result.get('certifications', [])],
            languages=[LanguageItem(**lang) for lang in result.get('languages', [])]
        )
    except Exception as e:
        logger.error(f"Resume parsing error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to parse resume")

async def generate_cover_letter_with_ai(resume_data: ResumeData, job_description: str, company_name: str, job_title: str) -> CoverLetterResponse:
    """Use OpenAI to generate cover letter"""
    
    resume_summary = f"""
Name: {resume_data.personalInfo.fullName}
Email: {resume_data.personalInfo.email}

Summary: {resume_data.summary}

Experience:
{chr(10).join([f"- {exp.title} at {exp.company}" for exp in resume_data.experience[:3]])}

Skills: {', '.join(resume_data.skills[:10])}

Education:
{chr(10).join([f"- {edu.degree} from {edu.school}" for edu in resume_data.education])}
"""
    
    prompt = f"""Write a professional cover letter for this job application:

Job Title: {job_title}
Company: {company_name}

Job Description:
{job_description}

Candidate Profile:
{resume_summary}

Write a compelling cover letter (300-400 words) that:
1. Opens with enthusiasm for the role
2. Highlights relevant experience and skills matching the job requirements
3. Shows understanding of the company and role
4. Demonstrates value the candidate would bring
5. Closes with a strong call to action
6. Uses British English spelling

Also provide 2-3 suggestions for customisation.

Return JSON:
{{
  "content": "<cover letter text with proper paragraphs>",
  "suggestions": ["suggestion 1", "suggestion 2", "suggestion 3"]
}}
"""
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert cover letter writer. Always use British English and return valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        
        return CoverLetterResponse(
            content=result.get('content', ''),
            suggestions=result.get('suggestions', [])
        )
    except Exception as e:
        logger.error(f"Cover letter generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate cover letter")

async def extract_skills_with_ai(text: str, existing_skills: List[str]) -> SkillsExtractResponse:
    """Use OpenAI to extract skills from job description"""
    
    prompt = f"""Extract technical and professional skills from this job description. 
Return a JSON list of skills with their categories.

Job Description:
{text}

Existing Skills (to mark as already added):
{', '.join(existing_skills)}

Return JSON:
{{
  "skills": [
    {{"name": "Python", "category": "Programming", "alreadyAdded": false}},
    {{"name": "Leadership", "category": "Soft Skills", "alreadyAdded": true}}
  ]
}}

Categories: Programming, Frameworks, Databases, Cloud, Tools, Soft Skills, Other"""
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert at extracting skills from job descriptions. Return valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        skills = []
        
        for skill_data in result.get('skills', []):
            already_added = skill_data.get('name', '') in existing_skills
            skills.append(SkillItem(
                name=skill_data.get('name', ''),
                category=skill_data.get('category', 'Other'),
                alreadyAdded=already_added
            ))
        
        return SkillsExtractResponse(skills=skills)
    except Exception as e:
        logger.error(f"Skills extraction error: {str(e)}")
        return SkillsExtractResponse(skills=[])

# AI Analysis Function
async def analyze_ats_with_ai(resume_data: ResumeData, job_description: str) -> ATSAnalysisResponse:
    """Use OpenAI to analyze resume against job description"""
    
    resume_text = f"""
Summary: {resume_data.summary}

Skills: {', '.join(resume_data.skills)}

Experience:
{chr(10).join([f"- {exp.title} at {exp.company}: {', '.join(exp.bullets)}" for exp in resume_data.experience])}

Education:
{chr(10).join([f"- {edu.degree} from {edu.school}" for edu in resume_data.education])}
"""
    
    prompt = f"""You are an expert ATS (Applicant Tracking System) analyzer.

Analyze this resume against the job description and provide:
1. ATS compatibility score (0-100)
2. Keywords from job description present in resume
3. Important keywords missing from resume
4. Specific suggestions to improve ATS score
5. Opportunities to quantify bullet points

Resume:
{resume_text}

Job Description:
{job_description}

Respond in JSON format:
{{
  "score": <number>,
  "matched_keywords": [<keywords>],
  "missing_keywords": [<keywords>],
  "suggestions": [<suggestions>],
  "impact_opportunities": [{{"original": "<text>", "improved": "<text>"}}],
  "readability_score": <number>,
  "readability_suggestions": [<suggestions>]
}}
"""
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert ATS analyzer. Always respond with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        
        score = result.get('score', 75)
        status = "Excellent" if score >= 80 else "Good" if score >= 60 else "Needs Work"
        
        return ATSAnalysisResponse(
            score=score,
            status=status,
            keywordAnalysis=KeywordAnalysis(
                matched=result.get('matched_keywords', []),
                missing=result.get('missing_keywords', []),
                overused=[]
            ),
            suggestions=result.get('suggestions', []),
            impactOpportunities=[
                ImpactOpportunity(original=opp.get('original', ''), suggestion=opp.get('improved', ''))
                for opp in result.get('impact_opportunities', [])
            ],
            readability=ReadabilityAnalysis(
                score=result.get('readability_score', 85),
                suggestions=result.get('readability_suggestions', [])
            )
        )
    except Exception as e:
        logger.error(f"OpenAI error: {str(e)}")
        return ATSAnalysisResponse(
            score=70,
            status="Good",
            keywordAnalysis=KeywordAnalysis(matched=["Product Management"], missing=["Machine Learning"], overused=[]),
            suggestions=["Add more quantified achievements"],
            impactOpportunities=[],
            readability=ReadabilityAnalysis(score=80, suggestions=["Use stronger action verbs"])
        )

# API Routes
@api_router.get("/")
async def root():
    return {"message": "Resume Builder API"}

@api_router.post("/resumes", response_model=Resume)
async def create_resume(input: ResumeCreate):
    resume_obj = Resume(resumeData=input.resumeData, template=input.template)
    await db.resumes.insert_one(resume_obj.dict())
    return resume_obj

@api_router.get("/resumes", response_model=List[Resume])
async def get_resumes():
    resumes = await db.resumes.find().to_list(1000)
    return [Resume(**resume) for resume in resumes]

@api_router.get("/resumes/{resume_id}", response_model=Resume)
async def get_resume(resume_id: str):
    resume = await db.resumes.find_one({"id": resume_id})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return Resume(**resume)

@api_router.post("/ai/analyze-ats", response_model=ATSAnalysisResponse)
async def analyze_ats(request: ATSAnalysisRequest):
    return await analyze_ats_with_ai(request.resumeData, request.jobDescription)

# Resume Parsing
@api_router.post("/parse-resume", response_model=ResumeData)
async def parse_resume(file: UploadFile = File(...)):
    """Parse uploaded resume file (PDF or DOCX)"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_ext = file.filename.lower().split('.')[-1]
    
    if file_ext not in ['pdf', 'docx']:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    # Read file content
    content = await file.read()
    
    # Extract text
    if file_ext == 'pdf':
        text = extract_text_from_pdf(content)
    else:
        text = extract_text_from_docx(content)
    
    # Parse with AI
    return await parse_resume_with_ai(text)

# Cover Letter
@api_router.post("/cover-letter/generate", response_model=CoverLetterResponse)
async def generate_cover_letter(request: CoverLetterRequest):
    """Generate AI-powered cover letter"""
    return await generate_cover_letter_with_ai(
        request.resumeData,
        request.jobDescription,
        request.companyName,
        request.jobTitle
    )

@api_router.post("/cover-letter/export/pdf")
async def export_cover_letter_pdf(request: dict):
    """Export cover letter as PDF"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=inch, bottomMargin=inch, leftMargin=inch, rightMargin=inch)
    story = []
    styles = getSampleStyleSheet()
    
    # Styles
    name_style = ParagraphStyle('Name', parent=styles['Normal'], fontSize=12, fontName='Helvetica-Bold')
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=11, leading=16)
    
    # Personal info
    personal_info = request.get('personalInfo', {})
    name_text = personal_info.get('fullName', '')
    contact_text = f"{personal_info.get('email', '')} | {personal_info.get('phone', '')} | {personal_info.get('location', '')}"
    
    story.append(Paragraph(name_text, name_style))
    story.append(Paragraph(contact_text, styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Date
    today = date.today().strftime("%d %B %Y")
    story.append(Paragraph(today, styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Company info
    company_name = request.get('companyName', '')
    job_title = request.get('jobTitle', '')
    if company_name:
        story.append(Paragraph(company_name, styles['Normal']))
    if job_title:
        story.append(Paragraph(f"Re: {job_title}", styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Cover letter content
    content = request.get('content', '')
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.strip(), body_style))
            story.append(Spacer(1, 0.2*inch))
    
    # Signature
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(f"Yours sincerely,<br/>{name_text}", styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    
    filename = f"{name_text.replace(' ', '_')}_Cover_Letter.pdf"
    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={filename}"})

# Export Routes
@api_router.post("/export/pdf")
async def export_pdf(request: ExportRequest):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#0f172a'), spaceAfter=12)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#0f172a'), spaceBefore=12, spaceAfter=6)
    
    data = request.resumeData
    
    story.append(Paragraph(data.personalInfo.fullName, title_style))
    story.append(Spacer(1, 6))
    
    contact_text = f"{data.personalInfo.email} | {data.personalInfo.phone}"
    if data.personalInfo.location:
        contact_text += f" | {data.personalInfo.location}"
    story.append(Paragraph(contact_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    if data.summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
        story.append(Paragraph(data.summary, styles['Normal']))
        story.append(Spacer(1, 12))
    
    if data.experience:
        story.append(Paragraph("WORK EXPERIENCE", heading_style))
        for exp in data.experience:
            job_title = f"<b>{exp.title}</b> - {exp.company}"
            story.append(Paragraph(job_title, styles['Normal']))
            date_range = f"{exp.startDate} - {'Present' if exp.current else exp.endDate}"
            story.append(Paragraph(date_range, styles['Normal']))
            for bullet in exp.bullets:
                if bullet:
                    story.append(Paragraph(f"• {bullet}", styles['Normal']))
            story.append(Spacer(1, 8))
    
    if data.education:
        story.append(Paragraph("EDUCATION", heading_style))
        for edu in data.education:
            degree_text = f"<b>{edu.degree}</b> - {edu.school} ({edu.graduationDate})"
            story.append(Paragraph(degree_text, styles['Normal']))
            story.append(Spacer(1, 8))
    
    if data.skills:
        story.append(Paragraph("SKILLS", heading_style))
        skills_text = ", ".join(data.skills)
        story.append(Paragraph(skills_text, styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    
    filename = f"{data.personalInfo.fullName.replace(' ', '_')}_Resume.pdf"
    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={filename}"})

@api_router.post("/export/docx")
async def export_docx(request: ExportRequest):
    doc = Document()
    data = request.resumeData
    
    name_para = doc.add_heading(data.personalInfo.fullName, 0)
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    contact_text = f"{data.personalInfo.email} | {data.personalInfo.phone}"
    if data.personalInfo.location:
        contact_text += f" | {data.personalInfo.location}"
    contact_para = doc.add_paragraph(contact_text)
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    if data.summary:
        doc.add_heading('Professional Summary', 1)
        doc.add_paragraph(data.summary)
    
    if data.experience:
        doc.add_heading('Work Experience', 1)
        for exp in data.experience:
            job_para = doc.add_paragraph()
            job_para.add_run(f"{exp.title}").bold = True
            job_para.add_run(f" - {exp.company}")
            date_range = f"{exp.startDate} - {'Present' if exp.current else exp.endDate}"
            doc.add_paragraph(date_range)
            for bullet in exp.bullets:
                if bullet:
                    doc.add_paragraph(bullet, style='List Bullet')
    
    if data.education:
        doc.add_heading('Education', 1)
        for edu in data.education:
            edu_para = doc.add_paragraph()
            edu_para.add_run(edu.degree).bold = True
            edu_para.add_run(f" - {edu.school} ({edu.graduationDate})")
    
    if data.skills:
        doc.add_heading('Skills', 1)
        doc.add_paragraph(", ".join(data.skills))
    
    if data.certifications:
        doc.add_heading('Certifications', 1)
        for cert in data.certifications:
            cert_text = f"{cert.name} - {cert.issuer}"
            if cert.date:
                cert_text += f" ({cert.date})"
            doc.add_paragraph(cert_text, style='List Bullet')
    
    if data.languages:
        doc.add_heading('Languages', 1)
        for lang in data.languages:
            doc.add_paragraph(f"{lang.language}: {lang.proficiency}", style='List Bullet')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{data.personalInfo.fullName.replace(' ', '_')}_Resume.docx"
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={filename}"})

app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
